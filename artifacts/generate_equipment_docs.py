"""
Contoso Electronics - Equipment Knowledge Base Generator
=========================================================

Generates 15 equipment reference documents in both Word (.docx) and PDF (.pdf)
formats. These documents are used as enterprise knowledge for the Copilot Studio
maintenance agent described in docs/demo_proposal.md.

Run:
    python artifacts/generate_equipment_docs.py

Requires:
    pip install python-docx reportlab
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
)

COMPANY = "Contoso Electronics"
BRAND_COLOR = RGBColor(0x00, 0x4B, 0x8D)  # Contoso blue for Word
BRAND_COLOR_PDF = colors.HexColor("#004B8D")
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


# ---------------------------------------------------------------------------
# Equipment data
# ---------------------------------------------------------------------------
# Each equipment entry contains structured sections used to build a rich
# maintenance reference document.

EQUIPMENT = [
    {
        "name": "Digital Storage Oscilloscope",
        "asset_id": "CE-OSC-1200",
        "model": "Contoso WaveView DSO-4000X",
        "category": "Test & Measurement",
        "location": "Electrical Test Lab - Bench 3",
        "manufacturer": "Contoso Instruments Division",
        "install_date": "2023-03-14",
        "warranty_expiry": "2027-03-14",
        "overview": (
            "The Contoso WaveView DSO-4000X is a 4-channel digital storage "
            "oscilloscope used for capturing, visualizing, and analyzing "
            "electrical signals during board bring-up, quality assurance, and "
            "failure analysis. It supports advanced triggering, protocol "
            "decoding, and automated measurements for high-speed digital and "
            "analog designs."
        ),
        "specs": [
            ("Bandwidth", "1 GHz"),
            ("Channels", "4 analog + 16 digital (MSO)"),
            ("Sample Rate", "5 GSa/s per channel"),
            ("Memory Depth", "50 Mpts"),
            ("Vertical Resolution", "8-bit (12-bit high-res mode)"),
            ("Display", "12.1 in capacitive touchscreen"),
            ("Input Impedance", "1 MΩ / 50 Ω selectable"),
            ("Power", "100-240 VAC, 50/60 Hz, 200 W"),
        ],
        "maintenance": [
            ("Daily", "Wipe screen and inspect probes for damage."),
            ("Monthly", "Verify self-calibration routine completes without error."),
            ("Quarterly", "Clean cooling fan intake and check firmware version."),
            ("Annual", "Full metrology calibration by accredited lab (traceable)."),
        ],
        "safety": [
            "Never exceed the maximum rated input voltage on any channel.",
            "Use only properly rated probes and attenuators.",
            "Ensure the instrument is properly grounded before use.",
            "Do not operate with the enclosure removed.",
        ],
        "issues": [
            ("No trace on screen", "Check trigger settings and channel enable; verify probe attenuation ratio."),
            ("Self-cal fails", "Allow 20-minute warm-up, disconnect all inputs, retry calibration."),
            ("Distorted waveform", "Inspect probe grounding and compensation; replace worn probe tips."),
            ("Touchscreen unresponsive", "Reboot instrument; if persistent, log a work order for panel service."),
        ],
        "support": "test.equipment.support@contoso.com | ext. 4210",
    },
    {
        "name": "Soldering Station",
        "asset_id": "CE-SOL-0450",
        "model": "Contoso ThermaPro SS-90D",
        "category": "Assembly & Rework",
        "location": "SMT Rework Cell - Station 7",
        "manufacturer": "Contoso Assembly Tools",
        "install_date": "2024-01-09",
        "warranty_expiry": "2026-01-09",
        "overview": (
            "The Contoso ThermaPro SS-90D is a temperature-controlled soldering "
            "station designed for precision hand soldering and rework of "
            "surface-mount and through-hole components. It features rapid tip "
            "heat-up, closed-loop temperature control, and ESD-safe construction."
        ),
        "specs": [
            ("Power Output", "90 W"),
            ("Temperature Range", "100 °C - 450 °C"),
            ("Temperature Stability", "±2 °C"),
            ("Heat-up Time", "< 10 seconds to 350 °C"),
            ("Tip Type", "Contoso T12-series cartridge"),
            ("ESD Rating", "ESD-safe (< 2 Ω to ground)"),
            ("Power", "120/230 VAC, 50/60 Hz"),
        ],
        "maintenance": [
            ("Each shift", "Clean tip with brass sponge; re-tin before storage."),
            ("Weekly", "Inspect tip for pitting; check ground continuity."),
            ("Monthly", "Verify temperature accuracy with tip thermometer."),
            ("As needed", "Replace worn cartridge tips and heating element."),
        ],
        "safety": [
            "Tip temperatures exceed 400 °C - never touch the tip.",
            "Return the iron to its stand when not in use.",
            "Work in a ventilated area or use fume extraction.",
            "Wear ESD wrist strap when handling sensitive components.",
        ],
        "issues": [
            ("Tip not heating", "Check cartridge seating and station fuse; test with spare tip."),
            ("Poor solder wetting", "Clean and re-tin tip; verify temperature setpoint and flux."),
            ("Temperature drift", "Recalibrate; inspect for damaged thermocouple."),
            ("Error E-04 on display", "Heating element open circuit - replace cartridge."),
        ],
        "support": "assembly.support@contoso.com | ext. 4155",
    },
    {
        "name": "CO2 Laser Cutter",
        "asset_id": "CE-LAS-3300",
        "model": "Contoso PrecisionCut LC-150",
        "category": "Fabrication",
        "location": "Enclosure Fabrication Shop - Bay 2",
        "manufacturer": "Contoso Fabrication Systems",
        "install_date": "2022-08-22",
        "warranty_expiry": "2025-08-22",
        "overview": (
            "The Contoso PrecisionCut LC-150 is a CO2 laser cutting and "
            "engraving system used for fabricating enclosures, panels, gaskets, "
            "and acrylic fixtures. It combines a sealed CO2 laser source with a "
            "motion gantry and integrated fume extraction for clean, precise cuts."
        ),
        "specs": [
            ("Laser Type", "Sealed CO2, 150 W"),
            ("Work Area", "1300 x 900 mm"),
            ("Cutting Speed", "Up to 1000 mm/s"),
            ("Positioning Accuracy", "±0.02 mm"),
            ("Cooling", "Closed-loop water chiller"),
            ("Assist Gas", "Compressed air / nitrogen"),
            ("Power", "230 VAC, 50/60 Hz, 3.5 kW"),
        ],
        "maintenance": [
            ("Daily", "Clean lens and mirror; check chiller water level and temp."),
            ("Weekly", "Inspect belts and rails; lubricate linear guides."),
            ("Monthly", "Verify beam alignment; clean exhaust ducting."),
            ("Annual", "Replace chiller coolant; service laser tube if power drops."),
        ],
        "safety": [
            "Class 4 laser - never bypass safety interlocks or open lid during operation.",
            "Always run fume extraction; laser-cut materials release hazardous fumes.",
            "Keep a fire extinguisher nearby; never leave the machine unattended while cutting.",
            "Only cut approved materials - never cut PVC or chlorinated plastics.",
        ],
        "issues": [
            ("Incomplete cuts", "Clean lens, check focus height, verify laser power and speed."),
            ("Chiller alarm", "Check water level and ambient temperature; clear airflow."),
            ("Reduced engraving quality", "Realign mirrors; inspect lens for burns."),
            ("Axis not homing", "Check limit switches and gantry obstructions; inspect drive belt."),
        ],
        "support": "fab.support@contoso.com | ext. 4380",
    },
    {
        "name": "Reflow Soldering Oven",
        "asset_id": "CE-RFO-2100",
        "model": "Contoso ThermFlow RO-800",
        "category": "SMT Line",
        "location": "SMT Line 1 - Position 5",
        "manufacturer": "Contoso Assembly Tools",
        "install_date": "2021-11-30",
        "warranty_expiry": "2026-11-30",
        "overview": (
            "The Contoso ThermFlow RO-800 is an 8-zone convection reflow oven "
            "used to solder surface-mount assemblies by precisely following a "
            "thermal profile. It provides uniform heating, nitrogen inerting, and "
            "closed-loop conveyor control for high-yield SMT production."
        ),
        "specs": [
            ("Heating Zones", "8 top + 8 bottom"),
            ("Max Temperature", "350 °C per zone"),
            ("Conveyor Width", "50 - 400 mm adjustable"),
            ("Conveyor Speed", "0.1 - 2.0 m/min"),
            ("Atmosphere", "Air or Nitrogen"),
            ("Profile Storage", "200 recipes"),
            ("Power", "3-phase 400 VAC, 40 kW"),
        ],
        "maintenance": [
            ("Daily", "Check conveyor tracking and flux collection trays."),
            ("Weekly", "Clean flux from heater zones and exhaust; inspect chains."),
            ("Monthly", "Run profile verification with thermocouple board."),
            ("Quarterly", "Replace flux filters; inspect blower motors."),
        ],
        "safety": [
            "Surfaces and boards exit the oven extremely hot - use heat-resistant gloves.",
            "Ensure exhaust ventilation is operating before running production.",
            "Lockout/tagout before entering the oven tunnel for service.",
            "Nitrogen inerting can cause asphyxiation - ensure proper room ventilation.",
        ],
        "issues": [
            ("Cold solder joints", "Verify profile peak temperature and time-above-liquidus."),
            ("Conveyor stalls", "Check chain tension and drive motor; clear jams."),
            ("Zone not reaching setpoint", "Inspect heating element and SSR for that zone."),
            ("Excess flux residue", "Clean oven; check flux management filter saturation."),
        ],
        "support": "smt.line.support@contoso.com | ext. 4501",
    },
    {
        "name": "Pick and Place Machine",
        "asset_id": "CE-PNP-2200",
        "model": "Contoso RapidPlace PP-3200",
        "category": "SMT Line",
        "location": "SMT Line 1 - Position 3",
        "manufacturer": "Contoso Assembly Tools",
        "install_date": "2021-11-30",
        "warranty_expiry": "2026-11-30",
        "overview": (
            "The Contoso RapidPlace PP-3200 is a high-speed SMT pick-and-place "
            "machine that mounts surface-mount components onto printed circuit "
            "boards with vision-guided precision. It supports a wide component "
            "range from 01005 chips to large connectors and BGAs."
        ),
        "specs": [
            ("Placement Speed", "45,000 CPH (optimal)"),
            ("Placement Accuracy", "±25 µm @ 3σ"),
            ("Component Range", "01005 to 55 x 55 mm"),
            ("Feeder Capacity", "120 8mm feeders"),
            ("Vision System", "Dual on-head + fixed fiducial camera"),
            ("PCB Size", "50 x 50 to 510 x 460 mm"),
            ("Power", "3-phase 400 VAC, 6 kW"),
        ],
        "maintenance": [
            ("Daily", "Clean nozzles and check vacuum; verify feeder seating."),
            ("Weekly", "Lubricate XY gantry; clean camera lenses."),
            ("Monthly", "Run placement accuracy calibration; inspect nozzle wear."),
            ("Quarterly", "Replace worn nozzles and feeder springs; check belts."),
        ],
        "safety": [
            "Keep hands clear of the gantry - moving parts can cause injury.",
            "Never bypass light curtains or door interlocks.",
            "Follow ESD procedures when handling components and boards.",
            "Lockout/tagout before reaching into the machine envelope.",
        ],
        "issues": [
            ("Components not picked", "Check nozzle vacuum, feeder pitch, and part height in library."),
            ("Placement offset", "Recalibrate vision fiducials; verify PCB seating."),
            ("Frequent tombstoning", "Review paste deposit and placement force; check nozzle."),
            ("Feeder errors", "Reseat feeder; clean feeder contacts; replace if faulty."),
        ],
        "support": "smt.line.support@contoso.com | ext. 4501",
    },
    {
        "name": "Automated Optical Inspection System",
        "asset_id": "CE-AOI-2400",
        "model": "Contoso VisionCheck AOI-500",
        "category": "Quality Inspection",
        "location": "SMT Line 1 - Position 6",
        "manufacturer": "Contoso Vision Systems",
        "install_date": "2022-02-18",
        "warranty_expiry": "2027-02-18",
        "overview": (
            "The Contoso VisionCheck AOI-500 is an automated optical inspection "
            "system that detects assembly defects such as missing components, "
            "misalignment, solder bridges, and polarity errors immediately after "
            "reflow. It uses multi-angle lighting and AI-assisted defect "
            "classification to maximize detection while minimizing false calls."
        ),
        "specs": [
            ("Camera Resolution", "18 µm/pixel"),
            ("Lighting", "Multi-angle RGB + coaxial"),
            ("Inspection Speed", "Up to 55 cm²/s"),
            ("Min Component", "01005"),
            ("PCB Size", "50 x 50 to 510 x 460 mm"),
            ("Defect Library", "AI-assisted, self-learning"),
            ("Power", "230 VAC, 50/60 Hz, 1.2 kW"),
        ],
        "maintenance": [
            ("Daily", "Clean camera window; verify reference board passes."),
            ("Weekly", "Clean conveyor and lighting diffusers."),
            ("Monthly", "Recalibrate camera and lighting intensity."),
            ("Quarterly", "Update defect library; verify measurement gauge R&R."),
        ],
        "safety": [
            "Do not stare into inspection LEDs at close range.",
            "Follow ESD procedures when handling boards.",
            "Keep conveyor area clear during operation.",
            "Power down before cleaning camera optics.",
        ],
        "issues": [
            ("High false-call rate", "Retune thresholds; clean optics; verify lighting calibration."),
            ("Missed defects", "Update defect models; check camera focus and resolution."),
            ("Board not detected", "Verify conveyor sensors and fiducial recognition."),
            ("Inconsistent measurements", "Recalibrate; check for ambient light interference."),
        ],
        "support": "quality.support@contoso.com | ext. 4620",
    },
    {
        "name": "Wave Soldering Machine",
        "asset_id": "CE-WAV-2600",
        "model": "Contoso FlowSolder WS-600",
        "category": "Through-Hole Line",
        "location": "Through-Hole Line 2 - Position 4",
        "manufacturer": "Contoso Assembly Tools",
        "install_date": "2020-06-05",
        "warranty_expiry": "2025-06-05",
        "overview": (
            "The Contoso FlowSolder WS-600 is a wave soldering machine for "
            "through-hole and mixed-technology assemblies. It applies flux, "
            "preheats the board, and passes it over a molten solder wave to form "
            "reliable solder joints at production volumes."
        ),
        "specs": [
            ("Solder Pot Capacity", "250 kg"),
            ("Max Board Width", "400 mm"),
            ("Conveyor Speed", "0.5 - 2.0 m/min"),
            ("Preheat Zones", "3 (IR + convection)"),
            ("Flux System", "Spray fluxer"),
            ("Solder Alloy", "SAC305 lead-free"),
            ("Power", "3-phase 400 VAC, 30 kW"),
        ],
        "maintenance": [
            ("Daily", "Check solder level and dross; verify flux spray pattern."),
            ("Weekly", "Skim dross from solder pot; clean fingers/pallets."),
            ("Monthly", "Analyze solder alloy composition; clean flux nozzles."),
            ("Quarterly", "Inspect pump impeller and preheater elements."),
        ],
        "safety": [
            "Molten solder pot exceeds 260 °C - wear face shield and heat gloves when skimming.",
            "Flux fumes are hazardous - ensure extraction is active.",
            "Lead-free alloy is hot and heavy - use proper dross tools.",
            "Lockout/tagout before servicing the conveyor or pump.",
        ],
        "issues": [
            ("Solder skips/voids", "Check wave height, flux coverage, and preheat profile."),
            ("Excessive dross", "Reduce wave turbulence; verify nitrogen (if equipped)."),
            ("Bridging", "Adjust conveyor speed and exit angle; check board design."),
            ("Poor hole fill", "Increase preheat; verify flux activity and contact time."),
        ],
        "support": "through.hole.support@contoso.com | ext. 4540",
    },
    {
        "name": "Solder Paste Stencil Printer",
        "asset_id": "CE-SPP-2000",
        "model": "Contoso PrintPro SP-700",
        "category": "SMT Line",
        "location": "SMT Line 1 - Position 2",
        "manufacturer": "Contoso Assembly Tools",
        "install_date": "2021-11-30",
        "warranty_expiry": "2026-11-30",
        "overview": (
            "The Contoso PrintPro SP-700 is an automatic solder paste stencil "
            "printer that deposits precise volumes of solder paste onto PCB pads "
            "before component placement. It features vision alignment, "
            "programmable squeegee control, and automated stencil cleaning."
        ),
        "specs": [
            ("Print Area", "Up to 510 x 460 mm"),
            ("Alignment Accuracy", "±12.5 µm @ 6σ"),
            ("Print Force", "1 - 20 kg programmable"),
            ("Squeegee", "Metal blade, dual-direction"),
            ("Cycle Time", "< 8 seconds (excl. print stroke)"),
            ("Cleaning", "Automatic dry/wet/vacuum"),
            ("Power", "230 VAC, 50/60 Hz, 1.5 kW"),
        ],
        "maintenance": [
            ("Each shift", "Check paste roll condition and squeegee blades."),
            ("Daily", "Clean stencil underside; verify vision alignment."),
            ("Weekly", "Clean rails and support pins; check vacuum."),
            ("Monthly", "Recalibrate print pressure and camera; inspect blades."),
        ],
        "safety": [
            "Follow ESD procedures when loading boards.",
            "Solder paste contains hazardous metals - wear gloves and wash hands.",
            "Keep hands clear of moving squeegee head and print carriage.",
            "Use approved solvents in the automatic cleaner.",
        ],
        "issues": [
            ("Insufficient paste deposit", "Check aperture blockage; adjust print speed/pressure."),
            ("Smearing/bridging", "Clean stencil more frequently; verify gasketing and separation speed."),
            ("Misalignment", "Recalibrate fiducial cameras; verify board support."),
            ("Inconsistent volume", "Inspect squeegee wear; verify paste viscosity and roll size."),
        ],
        "support": "smt.line.support@contoso.com | ext. 4501",
    },
    {
        "name": "Function / Arbitrary Waveform Generator",
        "asset_id": "CE-FGN-1300",
        "model": "Contoso SignalGen FG-2200",
        "category": "Test & Measurement",
        "location": "Electrical Test Lab - Bench 5",
        "manufacturer": "Contoso Instruments Division",
        "install_date": "2023-05-20",
        "warranty_expiry": "2027-05-20",
        "overview": (
            "The Contoso SignalGen FG-2200 is a dual-channel function and "
            "arbitrary waveform generator used to stimulate circuits under test, "
            "generate clocks and modulation, and replay custom waveforms during "
            "design verification and production test."
        ),
        "specs": [
            ("Channels", "2 independent"),
            ("Max Frequency", "80 MHz (sine)"),
            ("Sample Rate", "500 MSa/s"),
            ("Arb Waveform Length", "16 Mpts/channel"),
            ("Amplitude Range", "1 mVpp - 10 Vpp (50 Ω)"),
            ("Modulation", "AM, FM, PM, FSK, PWM, sweep, burst"),
            ("Power", "100-240 VAC, 50/60 Hz"),
        ],
        "maintenance": [
            ("Monthly", "Verify output amplitude against calibrated meter."),
            ("Quarterly", "Update firmware; clean ventilation."),
            ("Annual", "Full calibration by accredited metrology lab."),
            ("As needed", "Inspect BNC connectors for wear."),
        ],
        "safety": [
            "Do not exceed rated output into low-impedance loads.",
            "Never apply external voltage to output terminals.",
            "Ensure proper grounding before connecting to DUT.",
            "Use correctly rated coaxial cables.",
        ],
        "issues": [
            ("No output", "Verify channel is enabled and output is turned on; check load setting."),
            ("Incorrect amplitude", "Check output load impedance setting (50 Ω vs High-Z)."),
            ("Distorted waveform", "Reduce amplitude for high frequency; check cable/termination."),
            ("Frequency error", "Allow warm-up; recalibrate timebase."),
        ],
        "support": "test.equipment.support@contoso.com | ext. 4210",
    },
    {
        "name": "Programmable DC Power Supply",
        "asset_id": "CE-PSU-1400",
        "model": "Contoso PowerBench PS-3060",
        "category": "Test & Measurement",
        "location": "Electrical Test Lab - Bench 2",
        "manufacturer": "Contoso Instruments Division",
        "install_date": "2023-04-11",
        "warranty_expiry": "2027-04-11",
        "overview": (
            "The Contoso PowerBench PS-3060 is a programmable triple-output DC "
            "power supply used to power boards under test with precise voltage "
            "and current control, sequencing, and protection features for safe, "
            "repeatable bench and production testing."
        ),
        "specs": [
            ("Outputs", "3 (2 x 0-30 V/5 A, 1 x 0-6 V/3 A)"),
            ("Line/Load Regulation", "< 0.01% + 3 mV"),
            ("Readback Accuracy", "0.05%"),
            ("Protection", "OVP, OCP, OTP"),
            ("Interface", "USB, LAN, GPIB"),
            ("Programming", "SCPI command set"),
            ("Power", "100-240 VAC, 50/60 Hz"),
        ],
        "maintenance": [
            ("Monthly", "Verify output voltage/current accuracy."),
            ("Quarterly", "Clean intake filter; check terminal tightness."),
            ("Annual", "Full calibration; verify protection trip points."),
            ("As needed", "Inspect binding posts and leads for damage."),
        ],
        "safety": [
            "Outputs can source high current - avoid short circuits.",
            "Verify polarity before connecting the DUT.",
            "Do not exceed rated wire gauge for output current.",
            "Enable OVP/OCP to protect sensitive devices under test.",
        ],
        "issues": [
            ("Output shuts off", "OVP/OCP tripped - check DUT and setpoints; reset protection."),
            ("Voltage not reaching setpoint", "In constant-current mode - check load current draw."),
            ("Inaccurate readback", "Recalibrate; verify remote sense connections."),
            ("No remote control", "Check interface cabling and SCPI address configuration."),
        ],
        "support": "test.equipment.support@contoso.com | ext. 4210",
    },
    {
        "name": "Bench Digital Multimeter",
        "asset_id": "CE-DMM-1100",
        "model": "Contoso PrecisionMeter DM-6500",
        "category": "Test & Measurement",
        "location": "Electrical Test Lab - Bench 1",
        "manufacturer": "Contoso Instruments Division",
        "install_date": "2023-02-01",
        "warranty_expiry": "2027-02-01",
        "overview": (
            "The Contoso PrecisionMeter DM-6500 is a 6½-digit bench digital "
            "multimeter used for precise measurement of voltage, current, "
            "resistance, frequency, and temperature during calibration, "
            "verification, and quality testing."
        ),
        "specs": [
            ("Display Digits", "6½ (1,200,000 counts)"),
            ("DCV Accuracy", "0.0035% of reading"),
            ("Max Voltage", "1000 V DC / 750 V AC"),
            ("Max Current", "10 A"),
            ("Functions", "V, I, Ω, freq, cap, temp, continuity"),
            ("Interface", "USB, LAN, GPIB"),
            ("Power", "100-240 VAC, 50/60 Hz"),
        ],
        "maintenance": [
            ("Monthly", "Verify against a known reference standard."),
            ("Quarterly", "Inspect test leads and fuses; clean terminals."),
            ("Annual", "Full traceable calibration."),
            ("As needed", "Replace current-input fuse after overload."),
        ],
        "safety": [
            "Never exceed input voltage/current ratings for the selected function.",
            "Use CAT-rated leads appropriate for the measurement.",
            "Disconnect before switching between current and voltage inputs.",
            "Replace fuses only with the specified type and rating.",
        ],
        "issues": [
            ("No current reading", "Check current-input fuse; verify lead in correct jack."),
            ("Unstable readings", "Check lead connections; enable appropriate filtering/NPLC."),
            ("Reading offset", "Perform zero/null; recalibrate if out of tolerance."),
            ("OL displayed", "Range/limit exceeded - select correct function and range."),
        ],
        "support": "test.equipment.support@contoso.com | ext. 4210",
    },
    {
        "name": "Environmental Test Chamber",
        "asset_id": "CE-ETC-3100",
        "model": "Contoso ClimaTest EC-2000",
        "category": "Reliability Lab",
        "location": "Reliability Lab - Chamber 1",
        "manufacturer": "Contoso Environmental Systems",
        "install_date": "2022-09-15",
        "warranty_expiry": "2027-09-15",
        "overview": (
            "The Contoso ClimaTest EC-2000 is a temperature and humidity test "
            "chamber used for environmental stress screening, thermal cycling, "
            "and reliability qualification of electronic assemblies to industry "
            "standards."
        ),
        "specs": [
            ("Temperature Range", "-70 °C to +180 °C"),
            ("Humidity Range", "10% to 98% RH"),
            ("Chamber Volume", "225 liters"),
            ("Temp Ramp Rate", "5 °C/min"),
            ("Uniformity", "±0.5 °C"),
            ("Controller", "Programmable, 100 profiles"),
            ("Power", "3-phase 400 VAC, 12 kW"),
        ],
        "maintenance": [
            ("Weekly", "Check water reservoir and drain; inspect door seal."),
            ("Monthly", "Clean condenser coils; verify temp/humidity accuracy."),
            ("Quarterly", "Inspect refrigeration system pressures."),
            ("Annual", "Full calibration; replace humidity sensor if drifting."),
        ],
        "safety": [
            "Chamber surfaces can be extremely hot or cold - use PPE.",
            "Do not open the door during extreme-temperature operation.",
            "Ensure adequate ventilation for refrigerant safety.",
            "Never place flammable or sealed pressurized items inside.",
        ],
        "issues": [
            ("Cannot reach low temp", "Check refrigerant charge and condenser airflow; verify door seal."),
            ("Humidity error", "Refill water reservoir; clean wick; verify sensor."),
            ("Temp overshoot", "Retune PID; reduce load mass; check sensor placement."),
            ("Frost buildup", "Run defrost cycle; inspect for door leaks."),
        ],
        "support": "reliability.support@contoso.com | ext. 4710",
    },
    {
        "name": "X-Ray Inspection System",
        "asset_id": "CE-XRI-3400",
        "model": "Contoso XView XR-900",
        "category": "Quality Inspection",
        "location": "Quality Lab - Room 4 (Shielded)",
        "manufacturer": "Contoso Vision Systems",
        "install_date": "2023-07-08",
        "warranty_expiry": "2028-07-08",
        "overview": (
            "The Contoso XView XR-900 is an X-ray inspection system used to "
            "examine hidden solder joints such as BGAs, QFNs, and internal "
            "structures for voids, bridging, and alignment defects that optical "
            "inspection cannot detect."
        ),
        "specs": [
            ("Tube Voltage", "Up to 130 kV"),
            ("Focal Spot", "1 µm (nanofocus)"),
            ("Magnification", "Up to 2000x geometric"),
            ("Detector", "Flat-panel digital"),
            ("Sample Manipulator", "5-axis, tilt to 70°"),
            ("Radiation Safety", "Fully shielded cabinet"),
            ("Power", "230 VAC, 50/60 Hz, 2 kW"),
        ],
        "maintenance": [
            ("Daily", "Verify interlocks and warning lamps function."),
            ("Weekly", "Clean sample stage; check detector calibration image."),
            ("Monthly", "Perform tube conditioning/warm-up cycle."),
            ("Annual", "Radiation survey by certified officer; detector calibration."),
        ],
        "safety": [
            "Ionizing radiation - never bypass door interlocks or shielding.",
            "Only trained, authorized operators may run the system.",
            "Wear personal dosimeter as required by radiation safety program.",
            "Report any interlock fault immediately and stop use.",
        ],
        "issues": [
            ("No X-ray image", "Verify tube warm-up complete; check interlocks and tube status."),
            ("Poor image contrast", "Adjust kV/current; run detector gain calibration."),
            ("Stage not moving", "Check manipulator drives and limit switches."),
            ("Interlock fault", "Do not override - inspect door switches; call service."),
        ],
        "support": "quality.support@contoso.com | ext. 4620",
    },
    {
        "name": "Conformal Coating Machine",
        "asset_id": "CE-CCM-2800",
        "model": "Contoso CoatShield CC-400",
        "category": "Finishing Line",
        "location": "Finishing Line 3 - Position 1",
        "manufacturer": "Contoso Finishing Systems",
        "install_date": "2022-04-27",
        "warranty_expiry": "2026-04-27",
        "overview": (
            "The Contoso CoatShield CC-400 is a selective conformal coating "
            "system that applies protective coating to PCBs to guard against "
            "moisture, dust, and chemical contamination. It uses programmable "
            "spray valves and vision for precise, masked-free selective coating."
        ),
        "specs": [
            ("Coating Type", "Acrylic, silicone, urethane"),
            ("Applicators", "Film-coat + needle valve"),
            ("Positioning Accuracy", "±0.1 mm"),
            ("Max PCB Size", "450 x 450 mm"),
            ("Coating Speed", "Up to 400 mm/s"),
            ("Curing", "Integrated UV / thermal option"),
            ("Power", "230 VAC, 50/60 Hz, 3 kW"),
        ],
        "maintenance": [
            ("Each shift", "Purge and clean spray valves; check material level."),
            ("Daily", "Inspect nozzles for clogging; verify spray pattern."),
            ("Weekly", "Clean coating booth and exhaust filters."),
            ("Monthly", "Recalibrate valves and vision; check UV lamp output."),
        ],
        "safety": [
            "Coating solvents are flammable - no open flames; ensure extraction.",
            "Wear gloves and respiratory protection as required by SDS.",
            "UV curing lamps emit hazardous UV - never look at the lamp.",
            "Store coating materials per SDS requirements.",
        ],
        "issues": [
            ("Uneven coating", "Clean/replace nozzle; verify viscosity and spray pressure."),
            ("Nozzle clogging", "Purge system; use recommended thinner; check material shelf life."),
            ("Coating in keep-out areas", "Recalibrate vision alignment; verify program pattern."),
            ("Incomplete cure", "Check UV lamp hours/output or oven temperature and time."),
        ],
        "support": "finishing.support@contoso.com | ext. 4660",
    },
    {
        "name": "ESD-Protected Assembly Workstation",
        "asset_id": "CE-ESD-0100",
        "model": "Contoso SafeBench ESD-200",
        "category": "Assembly & Rework",
        "location": "Assembly Floor - Bench Group A",
        "manufacturer": "Contoso Assembly Tools",
        "install_date": "2024-02-15",
        "warranty_expiry": "2029-02-15",
        "overview": (
            "The Contoso SafeBench ESD-200 is an ESD-protected assembly "
            "workstation providing a grounded work surface, wrist-strap "
            "monitoring, ionization, and task lighting to protect sensitive "
            "electronic components from electrostatic discharge damage during "
            "manual assembly and inspection."
        ),
        "specs": [
            ("Work Surface", "Dissipative laminate, 1500 x 750 mm"),
            ("Surface Resistance", "10^6 - 10^9 Ω"),
            ("Wrist Strap Monitor", "Continuous dual-conductor"),
            ("Ionizer", "Balanced DC, ±35 V offset"),
            ("Grounding", "Common point ground (CPG)"),
            ("Lighting", "LED, glare-free, 1000 lux"),
            ("Power", "120/230 VAC, 50/60 Hz"),
        ],
        "maintenance": [
            ("Each shift", "Test wrist strap with tester before starting work."),
            ("Daily", "Clean work surface with ESD-safe cleaner."),
            ("Weekly", "Verify surface-to-ground resistance; check ionizer balance."),
            ("Quarterly", "Calibrate ionizer; inspect ground cords and snaps."),
        ],
        "safety": [
            "Always wear and test the wrist strap before handling components.",
            "Keep non-ESD-safe materials (plastics, foam cups) off the bench.",
            "Do not modify or defeat the grounding system.",
            "Ensure the common point ground is connected to building ground.",
        ],
        "issues": [
            ("Wrist strap monitor alarm", "Check strap fit and cord; replace worn strap or cord."),
            ("High surface resistance", "Clean surface; verify ground cord connection to CPG."),
            ("Ionizer imbalance", "Clean emitter points; recalibrate ionizer offset."),
            ("Static damage reports", "Audit grounding, humidity, and handling procedures."),
        ],
        "support": "assembly.support@contoso.com | ext. 4155",
    },
]


# ---------------------------------------------------------------------------
# Word (.docx) generation
# ---------------------------------------------------------------------------
def add_heading(doc, text, size, color=None, bold=True, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def build_docx(item, path):
    doc = Document()

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run(COMPANY)
    trun.bold = True
    trun.font.size = Pt(22)
    trun.font.color.rgb = BRAND_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run("Equipment Reference & Maintenance Guide")
    srun.font.size = Pt(12)
    srun.italic = True

    add_heading(doc, item["name"], 18, BRAND_COLOR, space_after=2)
    model_p = doc.add_paragraph()
    mr = model_p.add_run(f"{item['model']}")
    mr.font.size = Pt(12)
    mr.bold = True

    # Identification table
    add_heading(doc, "Asset Identification", 13, BRAND_COLOR)
    id_rows = [
        ("Asset ID", item["asset_id"]),
        ("Category", item["category"]),
        ("Manufacturer", item["manufacturer"]),
        ("Location", item["location"]),
        ("Installation Date", item["install_date"]),
        ("Warranty Expiry", item["warranty_expiry"]),
        ("Support Contact", item["support"]),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in id_rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
        cells[0].paragraphs[0].runs[0].bold = True

    # Overview
    add_heading(doc, "Overview", 13, BRAND_COLOR)
    doc.add_paragraph(item["overview"])

    # Specifications
    add_heading(doc, "Technical Specifications", 13, BRAND_COLOR)
    st = doc.add_table(rows=1, cols=2)
    st.style = "Light Grid Accent 1"
    hdr = st.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Value"
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].paragraphs[0].runs[0].bold = True
    for k, v in item["specs"]:
        cells = st.add_row().cells
        cells[0].text = k
        cells[1].text = v

    # Maintenance schedule
    add_heading(doc, "Preventive Maintenance Schedule", 13, BRAND_COLOR)
    mt = doc.add_table(rows=1, cols=2)
    mt.style = "Light Grid Accent 1"
    mhdr = mt.rows[0].cells
    mhdr[0].text = "Interval"
    mhdr[1].text = "Task"
    mhdr[0].paragraphs[0].runs[0].bold = True
    mhdr[1].paragraphs[0].runs[0].bold = True
    for interval, task in item["maintenance"]:
        cells = mt.add_row().cells
        cells[0].text = interval
        cells[1].text = task

    # Safety
    add_heading(doc, "Safety Guidelines", 13, BRAND_COLOR)
    for s in item["safety"]:
        doc.add_paragraph(s, style="List Bullet")

    # Troubleshooting
    add_heading(doc, "Common Issues & Troubleshooting", 13, BRAND_COLOR)
    tt = doc.add_table(rows=1, cols=2)
    tt.style = "Light Grid Accent 1"
    thdr = tt.rows[0].cells
    thdr[0].text = "Symptom"
    thdr[1].text = "Recommended Action"
    thdr[0].paragraphs[0].runs[0].bold = True
    thdr[1].paragraphs[0].runs[0].bold = True
    for symptom, action in item["issues"]:
        cells = tt.add_row().cells
        cells[0].text = symptom
        cells[1].text = action

    # Footer note
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        f"{COMPANY} internal maintenance documentation. For service requests, "
        f"contact {item['support']} or submit a work order through the "
        f"maintenance portal."
    )
    nr.italic = True
    nr.font.size = Pt(9)

    doc.save(path)


# ---------------------------------------------------------------------------
# PDF (.pdf) generation
# ---------------------------------------------------------------------------
def build_pdf(item, path):
    styles = getSampleStyleSheet()

    company_style = ParagraphStyle(
        "Company", parent=styles["Title"], textColor=BRAND_COLOR_PDF,
        fontSize=22, alignment=TA_CENTER, spaceAfter=2,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"], fontSize=11, alignment=TA_CENTER,
        textColor=colors.grey, spaceAfter=12,
    )
    h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"], textColor=BRAND_COLOR_PDF, fontSize=16,
        spaceBefore=6, spaceAfter=2,
    )
    model_style = ParagraphStyle(
        "Model", parent=styles["Normal"], fontSize=12, spaceAfter=8,
    )
    section = ParagraphStyle(
        "Section", parent=styles["Heading2"], textColor=BRAND_COLOR_PDF,
        fontSize=13, spaceBefore=12, spaceAfter=4,
    )
    body = ParagraphStyle(
        "Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=6,
    )
    bullet = ParagraphStyle(
        "Bullet", parent=body, leftIndent=14, bulletIndent=4,
    )
    footer = ParagraphStyle(
        "Footer", parent=styles["Normal"], fontSize=8, textColor=colors.grey,
        spaceBefore=16,
    )

    story = []
    story.append(Paragraph(COMPANY, company_style))
    story.append(Paragraph("Equipment Reference & Maintenance Guide", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.2, color=BRAND_COLOR_PDF))
    story.append(Spacer(1, 8))
    story.append(Paragraph(item["name"], h1))
    story.append(Paragraph(item["model"], model_style))

    def make_table(data, col_widths):
        tbl = Table(data, colWidths=col_widths, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND_COLOR_PDF),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0B0B0")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF3F8")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return tbl

    # Identification
    story.append(Paragraph("Asset Identification", section))
    id_data = [["Field", "Detail"]]
    id_data += [
        ["Asset ID", item["asset_id"]],
        ["Category", item["category"]],
        ["Manufacturer", item["manufacturer"]],
        ["Location", item["location"]],
        ["Installation Date", item["install_date"]],
        ["Warranty Expiry", item["warranty_expiry"]],
        ["Support Contact", item["support"]],
    ]
    story.append(make_table(id_data, [1.8 * inch, 4.4 * inch]))

    # Overview
    story.append(Paragraph("Overview", section))
    story.append(Paragraph(item["overview"], body))

    # Specifications
    story.append(Paragraph("Technical Specifications", section))
    spec_data = [["Parameter", "Value"]] + [[k, v] for k, v in item["specs"]]
    story.append(make_table(spec_data, [2.6 * inch, 3.6 * inch]))

    # Maintenance
    story.append(Paragraph("Preventive Maintenance Schedule", section))
    maint_data = [["Interval", "Task"]] + [[i, t] for i, t in item["maintenance"]]
    story.append(make_table(maint_data, [1.4 * inch, 4.8 * inch]))

    # Safety
    story.append(Paragraph("Safety Guidelines", section))
    for s in item["safety"]:
        story.append(Paragraph(s, bullet, bulletText="\u2022"))

    # Troubleshooting
    story.append(Paragraph("Common Issues & Troubleshooting", section))
    issue_data = [["Symptom", "Recommended Action"]] + [[s, a] for s, a in item["issues"]]
    story.append(make_table(issue_data, [2.0 * inch, 4.2 * inch]))

    story.append(Paragraph(
        f"{COMPANY} internal maintenance documentation. For service requests, "
        f"contact {item['support']} or submit a work order through the "
        f"maintenance portal.",
        footer,
    ))

    doc = SimpleDocTemplate(
        path, pagesize=letter,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        title=f"{COMPANY} - {item['name']}",
        author=COMPANY,
    )
    doc.build(story)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def slugify(item):
    name = item["name"].lower()
    for ch in [" ", "/", "&", "-"]:
        name = name.replace(ch, "_")
    while "__" in name:
        name = name.replace("__", "_")
    return f"{item['asset_id']}_{name.strip('_')}"


# Assign a single output format per equipment so the knowledge base is a
# realistic mix of Word and PDF documents (not duplicate copies).
FORMAT_BY_ASSET = {
    "CE-OSC-1200": "pdf",
    "CE-SOL-0450": "docx",
    "CE-LAS-3300": "pdf",
    "CE-RFO-2100": "docx",
    "CE-PNP-2200": "pdf",
    "CE-AOI-2400": "docx",
    "CE-WAV-2600": "pdf",
    "CE-SPP-2000": "docx",
    "CE-FGN-1300": "pdf",
    "CE-PSU-1400": "docx",
    "CE-DMM-1100": "pdf",
    "CE-ETC-3100": "docx",
    "CE-XRI-3400": "pdf",
    "CE-CCM-2800": "docx",
    "CE-ESD-0100": "pdf",
}


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generated = []
    for item in EQUIPMENT:
        base = slugify(item)
        fmt = FORMAT_BY_ASSET.get(item["asset_id"], "pdf")
        if fmt == "docx":
            path = os.path.join(OUTPUT_DIR, base + ".docx")
            build_docx(item, path)
        else:
            path = os.path.join(OUTPUT_DIR, base + ".pdf")
            build_pdf(item, path)
        generated.append((base, fmt))
        print(f"Generated: {base}.{fmt}")

    docx_count = sum(1 for _, f in generated if f == "docx")
    pdf_count = sum(1 for _, f in generated if f == "pdf")
    print(f"\nDone. {len(generated)} equipment documents written to:\n  "
          f"{OUTPUT_DIR}\n  ({docx_count} Word, {pdf_count} PDF)")


if __name__ == "__main__":
    main()
